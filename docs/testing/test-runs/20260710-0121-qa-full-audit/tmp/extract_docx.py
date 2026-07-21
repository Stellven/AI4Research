from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


def iter_block_text(doc: Document) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        blocks.append(
            {
                "kind": "paragraph",
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": paragraph.text,
            }
        )
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        blocks.append({"kind": "table", "index": table_index, "rows": rows})
    for section_index, section in enumerate(doc.sections, start=1):
        for role, container in (("header", section.header), ("footer", section.footer)):
            blocks.append(
                {
                    "kind": role,
                    "section": section_index,
                    "paragraphs": [p.text for p in container.paragraphs],
                    "tables": [
                        [[cell.text for cell in row.cells] for row in table.rows]
                        for table in container.tables
                    ],
                }
            )
    return blocks


def extract_ooxml_text(path: Path) -> dict[str, list[str]]:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    result: dict[str, list[str]] = {}
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            texts = [node.text or "" for node in root.findall(".//w:t", namespaces)]
            if texts:
                result[name] = texts
    return result


def main() -> None:
    source = Path(sys.argv[1]).resolve()
    output = Path(sys.argv[2]).resolve()
    document = Document(source)
    payload = {
        "source": str(source),
        "core_properties": {
            "title": document.core_properties.title,
            "subject": document.core_properties.subject,
            "author": document.core_properties.author,
            "last_modified_by": document.core_properties.last_modified_by,
            "created": document.core_properties.created.isoformat()
            if document.core_properties.created
            else None,
            "modified": document.core_properties.modified.isoformat()
            if document.core_properties.modified
            else None,
        },
        "blocks": iter_block_text(document),
        "ooxml_text": extract_ooxml_text(source),
    }
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
